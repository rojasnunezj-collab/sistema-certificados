# ====================================================================
# --- BLOQUE 0: Imports y Dependencias de Python-Docx ---
# ====================================================================
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.oxml.simpletypes import ST_TwipsMeasure, Twips

# ====================================================================
# --- BLOQUE 1: Parche Técnico Librearías Temporales (TwipsMeasure) ---
# ====================================================================
# Patch docx
original_convert_from_xml = ST_TwipsMeasure.convert_from_xml

@classmethod
def patch_convert_from_xml(cls, str_value):
    try:
        return Twips(int(str_value))
    except ValueError:
        try:
            return Twips(int(float(str_value)))
        except:
            return original_convert_from_xml(str_value)

ST_TwipsMeasure.convert_from_xml = patch_convert_from_xml

# ====================================================================
# --- BLOQUE 2: Utilidades de Modificación y Estilos de Tabla ---
# ====================================================================
def set_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_table_margins(table, top=0, bottom=0, left=10, right=10):
    tblPr = table._tbl.tblPr
    tblCellMar = parse_xml(f'''
    <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:top w:w="{top}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tblCellMar>
    ''')
    tblPr.append(tblCellMar)

# ====================================================================
# --- BLOQUE 3: Lógica Principal de Inyección Documental ---
# ====================================================================
def inyectar_tabla_en_docx(doc_io, data_items):
    doc = Document(doc_io)
    target_paragraph = None
    for p in doc.paragraphs:
        if '[[TABLA_NOTAS]]' in p.text:
            target_paragraph = p
            break
            
    if target_paragraph:
        target_paragraph.text = target_paragraph.text.replace('[[TABLA_NOTAS]]', '')
        for p in doc.paragraphs[:5]:
            if "CERTIFICADO" in p.text.upper():
                p.paragraph_format.space_after = Pt(0)
        
        table = doc.add_table(rows=1, cols=7)
        try:
            table.style = 'Table Grid'
        except:
            set_borders(table)
            
        table.autofit = False
        table.allow_autofit = False
        set_table_margins(table, top=72, bottom=72, left=30, right=30)

        widths = [Inches(0.75), Inches(0.75), Inches(0.75), Inches(3.0), Inches(0.75), Inches(0.75), Inches(0.75)]
        for i, col in enumerate(table.columns):
            col.width = widths[i]
        
        encabezados = ['Fecha', 'Placa', 'N° Guía', 'Descripción', 'Cantidad', 'Medida', 'Peso']
        hdr_cells = table.rows[0].cells
        for i, nombre in enumerate(encabezados):
            cell = hdr_cells[i]
            cell.text = nombre
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, "70ad47")
            
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1
                run = p.runs[0] if p.runs else p.add_run(nombre)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
        
        for item in data_items:
            row_cells = table.add_row().cells
            vals = [
                str(item.get('fecha_origen', '')),
                str(item.get('placa_origen', '')),
                str(item.get('guia_origen', '')),
                str(item.get('desc', '')),
                str(item.get('cant', '')),
                str(item.get('um', '')).upper(),
                str(item.get('peso', ''))
            ]
            
            for idx, valor in enumerate(vals):
                cell = row_cells[idx]
                cell.text = valor
                cell.width = widths[idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1
                    run = p.runs[0] if p.runs else p.add_run(valor)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

        tbl, p = table._tbl, target_paragraph._p
        p.addnext(tbl)

    new_buffer = io.BytesIO()
    doc.save(new_buffer)
    return new_buffer.getvalue()
