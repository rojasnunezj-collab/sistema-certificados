import streamlit as st
import pandas as pd
import google.generativeai as genai
import json
import io
import base64
import re
import time
import os
from datetime import datetime, timedelta
from dotenv import load_dotenv
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.oxml.simpletypes
from docx.oxml.simpletypes import ST_TwipsMeasure, Twips
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

# ==========================================
# 0. PATCHES & SETUP
# ==========================================

# Monkey patch para corregir error de parsing de floats en Twips
original_convert_from_xml = ST_TwipsMeasure.convert_from_xml

@classmethod
def patch_convert_from_xml(cls, str_value):
    try:
        return Twips(int(str_value))
    except ValueError:
        try:
            return Twips(int(float(str_value)))
        except:
            return original_convert_from_xml(str_value)

ST_TwipsMeasure.convert_from_xml = patch_convert_from_xml

load_dotenv()

st.set_page_config(page_title="Sistema Certificados", layout="wide")

# ==========================================
# 1. CREDENCIALES Y CONSTANTES
# ==========================================
API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY and "GEMINI_API_KEY" in st.secrets:
    API_KEY = st.secrets["GEMINI_API_KEY"]

if not API_KEY:
    try:
        if "gcp_service_account" in st.secrets and "GEMINI_API_KEY" in st.secrets["gcp_service_account"]:
            API_KEY = st.secrets["gcp_service_account"]["GEMINI_API_KEY"]
    except: pass

if not API_KEY:
    st.error("🚨 ERROR: Falta API KEY. Configura el archivo .env o st.secrets.")
    st.stop()

# --- MODELO DINÁMICO GLOBAL (Solución 404 y 429) ---
# --- MODELO DINÁMICO GLOBAL (Solución 404 y 429) ---
# --- MODELO DINÁMICO GLOBAL (Solución 404 y 429) ---
# --- MODELO DINÁMICO GLOBAL (Solución 404 y 429) ---
def get_verified_model():
    """ 
    1. Busca dinámicamente modelos Gemini válidos.
    2. Si falla, activa el PROTOCOLO DE FUERZA BRUTA probando variantes de nombres manuales.
    """
    BLACKLIST = ['experimental', 'preview', 'beta', 'robotics', '2.0', '2.5', '8b', 'gemma']
    
    # --- FASE 1: Búsqueda Dinámica en la Cuenta ---
    try:
        genai.configure(api_key=API_KEY)
        all_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        candidatos = []
        for m in all_models:
            name_lower = m.lower()
            if any(bad in name_lower for bad in BLACKLIST): continue # Filtro de Basura
            if 'gemini' not in name_lower: continue                  # Solo Gemini (Multimodal)
            
            score = 0
            if "1.5" in name_lower: score += 10
            if "flash" in name_lower: score += 5
            candidatos.append((score, m))
        
        candidatos.sort(key=lambda x: x[0], reverse=True)
        
        # Prueba de Fuego Dinámica
        for _, name in candidatos:
            try:
                model = genai.GenerativeModel(name)
                model.generate_content("test", request_options={"timeout": 5})
                return model, name # ¡Éxito Dinámico!
            except: continue
            
    except Exception:
        pass # Si falla listar modelos, vamos a la Fase 2 directo

    # --- FASE 2: PROTOCOLO DE RESCATE (Fuerza Bruta) ---
    # Probamos variantes específicas de nombres que suelen funcionar
    hardcoded_candidates = [
        "gemini-1.5-flash",             # Nombre corto
        "models/gemini-1.5-flash",      # Nombre con prefijo
        "gemini-1.5-flash-latest",      # Versión latest
        "models/gemini-1.5-flash-latest",
        "gemini-1.5-flash-001",         # Versión específica 001
        "models/gemini-1.5-flash-001",
        "gemini-1.5-pro",               # Respaldo Pro
        "models/gemini-1.5-pro",
        "gemini-pro",                   # El clásico
        "models/gemini-pro"
    ]
    
    errores = []
    for nombre in hardcoded_candidates:
        try:
            model = genai.GenerativeModel(nombre)
            # ¡LA PRUEBA REAL! Si esto pasa, el modelo existe y funciona.
            model.generate_content("test", request_options={"timeout": 5})
            return model, f"{nombre} (Rescate)"
        except Exception as e:
            errores.append(f"{nombre}: {str(e)}")
            continue

    # Si llegamos aquí, nada funcionó. Devolvemos error visible.
    return genai.GenerativeModel("gemini-1.5-flash"), f"⚠️ ERROR CRÍTICO: Ningún modelo respondió. Detalles: {errores[0] if errores else 'Sin conexión'}"

model, nombre_activo = get_verified_model()
st.sidebar.success(f"🟢 Motor Verificado: {nombre_activo}")

# IDs Google
ID_SHEET_REPOSITORIO = "14As5bCpZi56V5Nq1DRs0xl6R1LuOXLvRRoV26nI50NU"
ID_SHEET_CONTROL = "14As5bCpZi56V5Nq1DRs0xl6R1LuOXLvRRoV26nI50NU" 
DRIVE_FOLDER_ID = os.getenv("DRIVE_FOLDER_ID")
if not DRIVE_FOLDER_ID and "DRIVE_FOLDER_ID" in st.secrets:
    DRIVE_FOLDER_ID = st.secrets["DRIVE_FOLDER_ID"]

PLANTILLAS = {
    "EPMI S.A.C.": {
        "Comercialización/Disposición Final": os.getenv("TEMPLATE_EPMI_ID", "1d09vmlBlW_4yjrrz5M1XM8WpCvzTI4f11pERDbxFvNE"),
        "Peligroso y No Peligroso": os.getenv("TEMPLATE_EPMI_PELIGROSO_ID", "1QqqVJ2vCiAjiKKGt_zEpaImUB-q3aRurSiXjMEU--eg")
    },
    "INECOVE S.A.C.": {
        "Comercialización/Disposición Final": os.getenv("TEMPLATE_INECOVE_ID", "1MPzCwxR538osP3_br4VrTDybplqpTBtB08Jo"),
        "Peligroso y No Peligroso": os.getenv("TEMPLATE_INECOVE_PELIGROSO_ID", "1W-HyVSivqug13gBRBclBuICAOSBUHm1WN5cnqtMQcZY")
    }
}

# ==========================================
# 2. CONEXIÓN GOOGLE DRIVE/SHEETS
# ==========================================
def obtener_servicios():
    scopes = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/spreadsheets']
    creds = None
    try:
        if "gcp_service_account" in st.secrets:
            # FIX: Auto-correcion de typos en secrets (ej: https=// -> https://)
            info = dict(st.secrets["gcp_service_account"])
            for k, v in info.items():
                if isinstance(v, str):
                    info[k] = v.replace("https=//", "https://")
            
            creds = service_account.Credentials.from_service_account_info(info, scopes=scopes)
    except Exception as e:
        st.error(f"⚠️ Error de Autenticación: Revisa el formato de tus credenciales en Secrets. {e}")
        return None, None

    if not creds:
        try:
            creds = service_account.Credentials.from_service_account_file('secretos.json', scopes=scopes)
        except: return None, None

    try:
        return build('drive', 'v3', credentials=creds), build('sheets', 'v4', credentials=creds)
    except Exception as e:
        st.error(f"⚠️ Error conectando con Google API: {e}")
        return None, None

def registrar_en_control(datos_fila):
    _, sheets = obtener_servicios()
    if not sheets: return False
    try:
        sheets.spreadsheets().values().append(
            spreadsheetId=ID_SHEET_CONTROL, range="'historial'!A:J",
            valueInputOption="USER_ENTERED", body={"values": [datos_fila]}
        ).execute()
        return True
    except: return False

def subir_a_drive(contenido_bytes, nombre_archivo):
    # Validacion estricta pre-conexion
    folder_id = DRIVE_FOLDER_ID
    if not folder_id and "DRIVE_FOLDER_ID" in st.secrets:
        folder_id = st.secrets["DRIVE_FOLDER_ID"]
        
    if not folder_id:
        st.error("🚨 ERROR: No se puede subir el archivo. Faltan credenciales de DRIVE_FOLDER_ID en Secrets.")
        return None

    drive, _ = obtener_servicios()
    if not drive: return None
    
    try:
        file_metadata = {'name': f"{nombre_archivo}.docx", 'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}
        
        # Uso OBLIGATORIO de parents
        file_metadata['parents'] = [folder_id]
        
        media = MediaIoBaseUpload(io.BytesIO(contenido_bytes), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
        file = drive.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        return file.get('webViewLink')
    except Exception as e:
        err_msg = str(e)
        st.error(f"Error subiendo a Drive: {err_msg}")
        if "storageQuotaExceeded" in err_msg or "403" in err_msg:
            st.warning("⚠️ CUOTA EXCEDIDA (403): Revisa que 'DRIVE_FOLDER_ID' sea de una carpeta compartida con permiso de edición, no la raíz.")
        return None

# ==========================================
# 3. FORMATOS Y UTILIDADES
# ==========================================
def limpiar_monto(valor):
    """
    Convierte string a float.
    Maneja formato europeo/latino intercambiando comas por puntos para
    usar la lógica de split (mantener último punto como decimal).
    Ej: 3.700,00 -> 3.700.00 -> 3700.00
    Ej: 3,700.00 -> 3.700.00 -> 3700.00
    """
    if not valor: return 0.0
    s = str(valor).strip()
    
    # 1. Unificar separadores: todo a puntos
    s = s.replace(',', '.')
    
    # 2. Manejo de multiples puntos (ej: 3.580.00 o 3.700.00)
    # Si hay más de un punto, asumimos que solo el último es decimal
    if s.count('.') > 1:
        parts = s.split('.')
        # Unir todo menos el último con nada, y pegar el último con punto
        s = "".join(parts[:-1]) + '.' + parts[-1]
    
    # 3. Limpieza final caracteres no numéricos
    s = re.sub(r'[^\d.]', '', s)
    try:
        return float(s)
    except:
        return 0.0

def formato_inteligente(valor):
    """
    Formatea números:
    100.0 -> "100" (Sin decimales)
    3580.50 -> "3580.5" (Decimales justos)
    """
    try:
        f = float(valor)
        if f.is_integer():
            return f"{int(f)}"
        else:
            return f"{f}"
    except:
        return str(valor)

def obtener_fin_de_mes(fecha_str):
    try:
        dt = datetime.strptime(fecha_str, "%d/%m/%Y")
        next_month = dt.replace(day=28) + timedelta(days=4)
        res = next_month - timedelta(days=next_month.day)
        return res.strftime("%d/%m/%Y")
    except: return fecha_str

def limpiar_descripcion(texto):
    if not texto: return ""
    return re.sub(r'VEN\s*-\s*AMB\s*-\s*', '', str(texto).strip(), flags=re.IGNORECASE).strip()

def formato_nompropio(texto):
    return str(texto).strip().title() if texto else ""

def normalizar_fecha(fecha_str):
    if not fecha_str: return datetime.now().strftime("%d/%m/%Y")
    for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"]:
        try: return datetime.strptime(fecha_str.strip(), fmt).strftime("%d/%m/%Y")
        except: continue
    return fecha_str 

def formatear_guia(serie_str):
    if not serie_str or '-' not in str(serie_str): return serie_str
    try:
        p = str(serie_str).split('-')
        if len(p) == 2: return f"{p[0].strip()}-{str(int(p[1].strip()))}"
    except: pass
    return serie_str

@st.cache_data(show_spinner=False, ttl=10)
def leer_sheet_seguro(pestaña):
    _, s = obtener_servicios()
    if not s: return pd.DataFrame()
    try:
        r = s.spreadsheets().values().get(spreadsheetId=ID_SHEET_REPOSITORIO, range=f"'{pestaña}'!A1:Z1000").execute()
        v = r.get('values', [])
        if not v: return pd.DataFrame()
        return pd.DataFrame(v[1:], columns=v[0])
    except: return pd.DataFrame()

# ==========================================
# 4. LOGICA DOCX (INYECCIÓN DE TABLA)
# ==========================================
def set_borders(table):
    """Fallback para bordes manuales"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def set_cell_background(cell, color_hex):
    """Establece color de fondo de celda"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_table_margins(table, top=0, bottom=0, left=10, right=10):
    """Reduce márgenes internos de celdas"""
    tblPr = table._tbl.tblPr
    tblCellMar = parse_xml(f'''
    <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:top w:w="{top}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tblCellMar>
    ''')
    tblPr.append(tblCellMar)

def inyectar_tabla_en_docx(doc_io, data_items, servicio_global):
    """
    Recibe un BytesIO con el DOCX ya renderizado por docxtpl.
    Busca el marcador [[TABLA_NOTAS]] y lo reemplaza por una tabla real usando python-docx.
    """
    doc = Document(doc_io)
    
    # Buscar el párrafo con el marcador
    target_paragraph = None
    for p in doc.paragraphs:
        if '[[TABLA_NOTAS]]' in p.text:
            target_paragraph = p
            break
            
    if target_paragraph:
        # Limpiar el texto del marcador
        target_paragraph.text = target_paragraph.text.replace('[[TABLA_NOTAS]]', '')
        
        # FIX: Ajustar espaciado del parrafo del TITULO (asumiendo que esta antes o es general)
        # Iteramos los primeros parrafos para buscar el titulo y ajustar
        for p in doc.paragraphs[:5]:
            if "CERTIFICADO" in p.text.upper():
                p.paragraph_format.space_after = Pt(0)
        
        # Crear tabla y aplicar estilo/bordes
        table = doc.add_table(rows=1, cols=7)
        try:
            table.style = 'Table Grid'
        except:
            set_borders(table)
            
        # Ajustar ancho de tabla y columnas
        table.autofit = False
        table.allow_autofit = False
        
        # Padding de celdas: 0.05 pulgadas para equilibrio vertical
        # 0.05 inches * 1440 = 72 twips
        set_table_margins(table, top=72, bottom=72, left=30, right=30)

        # Anchos PROPORCIONALES (Total 7.5")
        # Fecha, Placa, Guia, Cant, Medida, Peso = 10% (0.75")
        # Descripcion = 40% (3.0")
        widths = [Inches(0.75), Inches(0.75), Inches(0.75), Inches(3.0), Inches(0.75), Inches(0.75), Inches(0.75)]
        
        for i, col in enumerate(table.columns):
            col.width = widths[i]
        
        # Encabezados
        encabezados = ['Fecha', 'Placa', 'N° Guía', 'Descripción', 'Cantidad', 'Medida', 'Peso']
        hdr_cells = table.rows[0].cells
        for i, nombre in enumerate(encabezados):
            cell = hdr_cells[i]
            cell.text = nombre
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Fondo Verde Estilo Sheets (#70ad47)
            set_cell_background(cell, "70ad47")
            
            # Centrar todos los párrafos
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # LIMPIEZA VERTICAL STRICTA
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1
                
                if p.runs:
                    run = p.runs[0]
                else:
                    run = p.add_run(nombre)
                    
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0) # Negro
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
        
        # Datos
        for item in data_items:
            row_cells = table.add_row().cells
            
            # Preparar valores (DIRECTOS del DataFrame ya formateado)
            # No re-procesamos para evitar errores de doble limpieza (Peso 0)
            p_cant = str(item.get('cant', ''))
            p_peso = str(item.get('peso', ''))
            
            vals = [
                str(item.get('fecha_origen', '')),
                str(item.get('placa_origen', '')),
                str(item.get('guia_origen', '')),
                str(item.get('desc', '')),
                p_cant,
                str(item.get('um', '')).upper(), # UM en Mayúsculas (DOBLE SEGURIDAD)
                p_peso 
            ]
            
            for idx, valor in enumerate(vals):
                cell = row_cells[idx]
                cell.text = valor
                cell.width = widths[idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Centrar todos los párrafos y quitar espacios
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # LIMPIEZA VERTICAL STRICTA (Space After 0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1
                    
                    if p.runs:
                        run = p.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                    else: 
                        run = p.add_run(valor) # Fallback
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)

        # Mover tabla después del párrafo
        tbl, p = table._tbl, target_paragraph._p
        p.addnext(tbl)

    new_buffer = io.BytesIO()
    doc.save(new_buffer)
    return new_buffer.getvalue()

# ==========================================
# 5. INTELIGENCIA ARTIFICIAL (GEMINI)
# ==========================================
def procesar_guia_ia(pdf_bytes):
    # EXTRACCIÓN SIN ERRORES (Usa el modelo global)
    # No re-configuramos aquí para ganar velocidad y usar la conexión global
    
    prompt = """
    Extrae en formato JSON: Correlativo, Fecha, RUC Remitente, RUC Destinatario, Placa, Chofer, Dirección de Llegada, Dirección de Partida, N° Guía y la Tabla de Pesos.
    REGLAS ESTRICTAS:
    - N° Guía: Extraer Serie-Numero completo.
    - Placa: Extraer placa del vehículo y carreta si existe.
    - Tabla: Extraer items con descripción completa.
    - PESOS (ANTI-ALUCINACIÓN): Si la unidad es 'UNID', 'UND', 'UNIDADES' y la guía NO tiene un peso explícito para ese ítem, DEVUELVE 0.00. NO CALCULES NI ESTIMES PESOS.
    - Para la Dirección de Partida, extrae la dirección principal (ej: Panamericana Sur Km 138.5). Luego, busca en OBSERVACIONES si existe un nombre de FUNDO o PLANTA. El resultado final debe ser: [Dirección Principal] - [Nombre del Fundo/Planta]. No omitas ninguna de las dos partes.
    
    JSON Esperado:
    {
        "fecha": "dd/mm/yyyy", 
        "serie": "T001-000000", 
        "vehiculo": "PLACA", 
        "punto_partida": "Dirección Concatenada", 
        "punto_llegada": "Dirección Completa", 
        "destinatario": "Razón Social", 
        "items": [
            {
                "desc": "Descripción del bien", 
                "cant": "Número", 
                "um": "Unidad (KG, UNID, GLN)", 
                "peso": "Peso (0.00 si es UNID y no explícito)"
            }
        ]
    }
    """
    
    if 'model' not in globals() or not model:
        st.error("Error: Modelo IA no inicializado globalmente.")
        return None

    try:
        time.sleep(1) 
        res = model.generate_content([prompt, {"mime_type": "application/pdf", "data": base64.b64encode(pdf_bytes).decode('utf-8')}])
        
        texto_limpio = res.text.replace("```json", "").replace("```", "")
        match = re.search(r'\{.*\}', texto_limpio, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        else:
            return None

    except Exception as e:
        st.error(f"Error IA: {e}")
        return None

# ==========================================
# 6. INTERFAZ STREAMLIT
# ==========================================
if 'ocr_data' not in st.session_state: st.session_state['ocr_data'] = None
if 'df_items' not in st.session_state: st.session_state['df_items'] = pd.DataFrame()
if 'datos_log_pendientes' not in st.session_state: st.session_state['datos_log_pendientes'] = {}
if 'generado' not in st.session_state: st.session_state['generado'] = False

with st.sidebar:
    st.header("Configuración")
    empresa_firma = st.selectbox("Empresa", list(PLANTILLAS.keys()))
    tipo_plantilla = st.selectbox("Plantilla", ["Comercialización/Disposición Final", "Peligroso y No Peligroso"])
    if st.button("Recargar"): st.cache_data.clear(); st.rerun()

st.title("Generador de Certificados")

if 'repo_data' not in st.session_state:
    st.session_state['repo_data'] = {
        "emisores": leer_sheet_seguro("EMPRESAS"),
        "clientes": leer_sheet_seguro("CLIENTES"),
        "servicios": leer_sheet_seguro("COMERCIALIZACION")
    }
repo = st.session_state['repo_data']

archivos = st.file_uploader("Subir Guías (PDF)", type=["pdf"], accept_multiple_files=True)

if archivos:
    if st.button("🔍 Procesar"):
        prog = st.progress(0)
        items, grl = [], None
        errores = 0
        total = len(archivos)
        
        for i, arc in enumerate(archivos):
            d = procesar_guia_ia(arc.read())
            if d:
                if not grl: grl = d 
                s, f, p = formatear_guia(d.get('serie','S/N')), d.get('fecha',''), d.get('vehiculo','')
                for it in d.get('items', []):
                    # Acumulando items de TODAS las guias
                    it.update({'guia_origen': s, 'fecha_origen': f, 'placa_origen': p})
                    items.append(it)
            else: errores += 1
            prog.progress((i+1)/total)
        
        time.sleep(0.5); prog.empty()
        
        if items:
            st.session_state['ocr_data'] = grl if grl else {}
            df = pd.DataFrame(items)
            for c in ['desc','cant','um','peso','fecha_origen','guia_origen','placa_origen']:
                if c not in df.columns: df[c] = ""
            
            df['peso'] = df['peso'].apply(lambda x: formato_inteligente(limpiar_monto(x)))
            df['cant'] = df['cant'].apply(lambda x: formato_inteligente(limpiar_monto(x)))
            df['desc'] = df['desc'].astype(str).str.upper()
            
            # REGLAS DE FORMATO: KG, GLN, UNID
            df['um'] = df['um'].apply(lambda x: 'KG' if 'KILO' in str(x).upper() else 'GLN' if 'GALO' in str(x).upper() else 'UNID' if 'UNIDA' in str(x).upper() else str(x).upper())
            
            df['desc'] = df['desc'].apply(limpiar_descripcion)
            df['fecha_origen'] = df['fecha_origen'].apply(normalizar_fecha)
            
            st.session_state['df_items'] = df
            st.success(f"✅ Procesado: {len(items)} items de {total} archivos.")
        else: st.error("❌ Falló: No se encontraron items.")

if st.session_state['ocr_data'] is not None:
    ocr = st.session_state['ocr_data']
    st.markdown("### Validación")
    
    # CSS Correlativo: Uso de aria-label para targeting preciso
    st.markdown('''
            <style>
            input[aria-label="Correlativo"] {
                background-color: #FFFF00 !important;
                color: black !important;
                font-weight: bold !important;
            }
            </style>
        ''', unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns(4)
    v_corr = c1.text_input("Correlativo", "001")
    fecha_base = normalizar_fecha(ocr.get('fecha'))
    cont_f = c2.container()
    
    opt_f = c2.radio("Regla Fecha:", ["COMERCIALIZACION (FIN DE MES)", "DISPOSICION FINAL +2"], label_visibility="collapsed")
    
    try:
        if "FIN DE MES" in opt_f: 
            f_calc = obtener_fin_de_mes(fecha_base)
            tipo_operacion_simple = "Comercialización"
        else: 
            f_calc = (datetime.strptime(fecha_base, "%d/%m/%Y")+timedelta(days=2)).strftime("%d/%m/%Y")
            tipo_operacion_simple = "Disposición Final"
    except: 
        f_calc = fecha_base
        tipo_operacion_simple = "Indefinido"

    v_fec_emis = cont_f.text_input("F. Emisión", value=f_calc)

    # Mostrar info del primer archivo como referencia
    v_guia_ref = formatear_guia(ocr.get('serie')) if len(archivos) == 1 else "VARIAS"
    v_placa_ref = ocr.get('vehiculo') if len(archivos) == 1 else "VARIAS"
    
    # FIX: Definir variables v_guia y v_placa para uso posterior
    v_guia = c3.text_input("Guía", v_guia_ref, disabled=True)
    v_placa = c4.text_input("Placa", v_placa_ref, disabled=True)


    # Sync Partida
    partida_base = formato_nompropio(ocr.get('punto_partida',''))
    if "v_partida" not in st.session_state: st.session_state["v_partida"] = partida_base
    v_partida = st.text_input("Partida", key="v_partida")

    # Sync Llegada
    llegada_base = formato_nompropio(ocr.get('punto_llegada',''))
    if "v_llegada" not in st.session_state: st.session_state["v_llegada"] = llegada_base
    v_llegada = st.text_input("Llegada", key="v_llegada")
    
    v_dest = st.text_input("Destinatario", ocr.get('destinatario',''), key="txt_destinatario")

    v_items = st.data_editor(st.session_state['df_items'], num_rows="dynamic", use_container_width=True,
        column_config={"guia_origen": st.column_config.TextColumn("Guía", disabled=True)}, key="editor_items")
    
    c_a, c_b = st.columns(2)
    with c_a:
        v_emi = st.selectbox("Emisor", repo['emisores']['EMPRESA'].unique() if not repo['emisores'].empty else [])
        v_ruc_e, v_reg_e = "", ""
        if not repo['emisores'].empty and v_emi:
            try:
                row_e = repo['emisores'][repo['emisores']['EMPRESA']==v_emi].iloc[0]
                v_ruc_e, v_reg_e = str(row_e['RUC']), str(row_e['REGISTRO'])
            except: pass
        st.caption(f"RUC: {v_ruc_e} | REG: {v_reg_e}")
        v_tit = st.selectbox("Título", repo['servicios'].iloc[:,0].unique() if not repo['servicios'].empty else [])

    with c_b:
        v_cli = st.selectbox("Cliente", repo['clientes']['EMPRESA'].unique() if not repo['clientes'].empty else [])
        v_ruc_c = ""
        if not repo['clientes'].empty and v_cli:
            try:
                row_c = repo['clientes'][repo['clientes']['EMPRESA']==v_cli].iloc[0]
                v_ruc_c = str(row_c['RUC'])
            except: pass
        st.caption(f"RUC: {v_ruc_c}")
        v_serv = st.selectbox("Servicio", repo['servicios'].iloc[:,1].unique() if not repo['servicios'].empty else [])
        v_res = st.selectbox("Residuo", repo['servicios'].iloc[:,2].unique() if not repo['servicios'].empty else [])

    dest_final = v_dest if "EPMI" not in str(v_dest).upper() else "EPMI S.A.C."

    st.divider()
    tab1, tab2 = st.tabs(["Generar", "Registrar"])

    with tab1:
        if st.button("GENERAR CERTIFICADO", type="primary"):
            drive, _ = obtener_servicios()
            if drive:
                try:
                    # 1. Descargar plantilla
                    id_p = PLANTILLAS[empresa_firma][tipo_plantilla]
                    req = drive.files().export_media(fileId=id_p, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    fh = io.BytesIO()
                    dl = MediaIoBaseDownload(fh, req)
                    done = False
                    while not done: _, done = dl.next_chunk()
                    
                    # 2. Renderizar Docxtpl
                    doc = DocxTemplate(io.BytesIO(fh.getvalue()))
                    ctx = {
                        "CORRELATIVO": v_corr, "TITULO": v_tit, "REGISTRO": v_reg_e,
                        "EMPRESA": v_emi, "RUC_EMPRESA": v_ruc_e, "RUC": v_ruc_e, 
                        "CLIENTE": v_cli, "RUC_CLIENTE": v_ruc_c, "RAZON_SOCIAL_CLIENTE": v_cli,
                        "SERVICIO_O_COMPRA": v_serv, "TIPO_DE_RESIDUO": v_res,
                        "PUNTO_PARTIDA": st.session_state["v_partida"], 
                        "DIRECCION_EMPRESA": st.session_state["v_llegada"], 
                        "DIRECCION_LLEGADA": st.session_state["v_llegada"], 
                        "LLEGADA": st.session_state["v_llegada"],
                        "EMPRESA_2": dest_final, "FECHA_EMISION": v_fec_emis,
                        "DESTINATARIO_FINAL": st.session_state["txt_destinatario"]
                    }
                    doc.render(ctx)
                    buf_tpl = io.BytesIO()
                    doc.save(buf_tpl)

                    # 3. Inyectar Tabla (Usando v_items editado por usuario)
                    # Procesar DataFrame antes de inyectar
                    df_final = v_items.copy()
                    df_final['um'] = df_final['um'].astype(str).str.upper()
                    df_final['cant'] = df_final['cant'].apply(lambda x: formato_inteligente(limpiar_monto(x)))
                    df_final['peso'] = df_final['peso'].apply(lambda x: formato_inteligente(limpiar_monto(x)))

                    items_para_tabla = df_final.to_dict('records')
                    final_bytes = inyectar_tabla_en_docx(io.BytesIO(buf_tpl.getvalue()), items_para_tabla, v_serv)
                    
                    st.session_state['word_buffer'] = final_bytes
                    
                    # Calcular Peso Total Seguro (Sobre items editados)
                    peso_t = sum([limpiar_monto(x) for x in df_final['peso']])
                    name_safe = f"{empresa_firma} - {tipo_operacion_simple} - {v_corr}".replace("/", "-")
                    st.session_state['nombre_archivo_final'] = name_safe
                    
                    # 4. Subida a Drive Manual (Desactivado Auto para evitar 403)
                    # link_drive = subir_a_drive(final_bytes, name_safe)
                    link_drive = None 
                    st.session_state['link_drive_generado'] = "" 

                    st.session_state['datos_log_pendientes'] = {
                        "fec_emis": v_fec_emis, "emi": v_emi, "tit": tipo_operacion_simple, 
                        "cli": v_cli, "ruc_c": v_ruc_c, "guia": v_guia, "res": v_res,
                        "cert_name": name_safe, "peso": peso_t              
                    }
                    
                    st.success(f"✅ Generado Correctamente: {name_safe}")
                    st.success(f"📍 Certificado generado con la dirección: {v_llegada}")
                    st.session_state['generado'] = True
                    st.rerun()

                except Exception as e: st.error(f"Error: {e}")
            else:
                st.error("No se pudo conectar con Google Drive.")

        if st.session_state.get('generado'):
            # Botones persistentes despues de generar
            fn = st.session_state.get('nombre_archivo_final', "Borrador")
            st.download_button("📩 Bajar Copia Local", st.session_state['word_buffer'], f"{fn}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.info("ℹ️ Descarga el archivo y súbelo manualmente a Drive.")
            st.link_button("📂 Abrir carpeta de Drive para subir archivo", "https://drive.google.com/drive/u/0/folders/1prb1KJZG_BYQSAicLWPP1t_nn58QVO--")

    with tab2:
        # Registro Manual: Campos vacíos para que el usuario pegue los links
        u_d = st.text_input("Link DOC:", value="")
        u_p = st.text_input("Link PDF:", value="")
        if st.button("🏁 Registrar"):
            if not st.session_state.get('datos_log_pendientes') or not u_d or not u_p:
                st.warning("⚠️ Faltan datos (Link Doc/PDF o generar primero)")
            else:
                d = st.session_state['datos_log_pendientes']
                f = [d['fec_emis'], d['emi'], d['tit'], d['cli'], d['ruc_c'], d['guia'], "FINALIZADO", d['cert_name'], u_d, u_p]
                if registrar_en_control(f): st.success("✅ Registrado en Sheets"); st.balloons()
